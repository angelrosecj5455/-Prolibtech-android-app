# # importing required modules
# import PyPDF2
#
# # creating a pdf file object
# pdfFileObj = open('C:\\Users\\Angel Rose\\PycharmProjects\\project_management\\static\\report\\20220429-104419.pdf', 'rb')
#
# # creating a pdf reader object
# pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
#
# # printing number of pages in pdf file
# print(pdfReader.numPages)
#
# # creating a page object
# ls=[]
# for i in range(0,pdfReader.numPages):
#     print(i)
#     pageObj = pdfReader.getPage(i)
#
#     # extracting text from page
#     print(pageObj.extractText())
#     ls.append(pageObj.extractText())
#
# # closing the pdf file object
# pdfFileObj.close()
#
# print("--------------------",ls)

from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import nltk
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))
import aspose.words as aw

# load the PDF file
doc = aw.Document("C:\\Users\\Angel Rose\\PycharmProjects\\project_management\\static\\report\\20220429-104419.pdf")

# convert PDF to Word DOCX format
doc.save("pdf-to-word.docx")
import docx
doc = docx.Document('C:\\Users\\Angel Rose\\PycharmProjects\\project_management\\pdf-to-word.docx')  # Creating word reader object.
data = ""
fullText = []
for para in doc.paragraphs:
    fullText.append(para.text)
    data = '\n'.join(fullText)

print(data)

data.replace(",", "").replace(".", "").replace("0", "").replace("1", "").replace("2", "").replace("3",
                                                                                                    "").replace(
        "4", "").replace("5", "").replace("6", "").replace("7", "").replace("8", "").replace("0", "").replace("!",
            "").replace("@", "").replace("'","").replace(",","")
ss = word_tokenize(data)
print(ss)
filtered_sentence = []

for w in ss:
    if w not in stop_words:
        filtered_sentence.append(w)

print(ss)
print(filtered_sentence)
# aa = list(session["wordslist"])
# print(aa)
aa = []
for i in filtered_sentence:
    aa.append(i)
# session["wordslist"] = aa
print(aa)

from collections import Counter

counts = Counter(aa)
print(counts)
k = counts.keys()
print(k)
v = counts.values()